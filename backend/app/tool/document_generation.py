import os
import json
import yaml
import pandas as pd
import markdown
from typing import Literal, Optional, Any
from pathlib import Path
from xhtml2pdf import pisa
from docx import Document
from app.tool.base import BaseTool, ToolResult
from app.logger import logger
import io

class UniversalFileGenerator(BaseTool):
    name: str = "generate_file"
    description: str = (
        "Generate various types of files including documents (PDF, DOCX), spreadsheets (Excel, CSV), "
        "and data files (JSON, YAML, XML). "
        "Use this tool whenever the user asks to create a file or report in a specific format. "
        "For documents (PDF/DOCX), provide Markdown content. "
        "For data/spreadsheets, provide JSON-formatted string representing the data."
    )
    parameters: dict = {
        "type": "object",
        "properties": {
            "content": {
                "type": "string",
                "description": "The content of the file. For PDF/DOCX, use Markdown. For JSON/YAML/XML/CSV/Excel, use a JSON string representing the data structure (list of dicts or dict).",
            },
            "filename": {
                "type": "string",
                "description": "The name of the file to save (e.g., 'report.pdf', 'data.xlsx').",
            },
            "format": {
                "type": "string",
                "enum": ["pdf", "docx", "json", "yaml", "xml", "csv", "xlsx"],
                "description": "The format of the file. If not provided, it will be inferred from the filename extension.",
            },
        },
        "required": ["content", "filename"],
    }

    async def execute(self, content: str, filename: str, format: Optional[str] = None) -> ToolResult:
        try:
            # Determine format from filename if not explicitly provided
            if not format:
                ext = os.path.splitext(filename)[1].lower().lstrip('.')
                if ext in ["pdf", "docx", "json", "yaml", "yml", "xml", "csv", "xlsx"]:
                    format = "yaml" if ext == "yml" else ext
                else:
                    return ToolResult(error=f"Could not infer supported format from filename '{filename}'. Please specify format.")

            # Ensure workspace directory exists
            workspace_root = Path(os.getcwd()) / "workspace"
            workspace_root.mkdir(parents=True, exist_ok=True)
            
            file_path = workspace_root / filename

            if format == "pdf":
                return self._generate_pdf(content, file_path)
            elif format == "docx":
                return self._generate_docx(content, file_path)
            elif format in ["json", "yaml", "xml", "csv", "xlsx"]:
                return self._generate_data_file(content, file_path, format)
            else:
                return ToolResult(error=f"Unsupported format: {format}")

        except Exception as e:
            logger.error(f"Error generating file: {e}")
            return ToolResult(error=f"Failed to generate file: {str(e)}")

    def _generate_pdf(self, content: str, file_path: Path) -> ToolResult:
        try:
            # Convert Markdown to HTML
            html_content = markdown.markdown(content, extensions=['extra', 'codehilite'])
            
            # Add basic styling
            styled_html = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Helvetica, sans-serif; font-size: 12pt; line-height: 1.5; }}
                    h1 {{ font-size: 24pt; color: #333333; border-bottom: 1px solid #cccccc; padding-bottom: 10px; }}
                    h2 {{ font-size: 18pt; color: #444444; margin-top: 20px; }}
                    h3 {{ font-size: 14pt; color: #555555; }}
                    code {{ background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
                    pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                    blockquote {{ border-left: 4px solid #cccccc; padding-left: 10px; color: #666666; }}
                    table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            with open(file_path, "wb") as pdf_file:
                pisa_status = pisa.CreatePDF(styled_html, dest=pdf_file)

            if pisa_status.err:
                return ToolResult(error=f"PDF generation error: {pisa_status.err}")
            
            return ToolResult(output=f"PDF document generated successfully: {file_path}")

        except Exception as e:
            raise e

    def _generate_docx(self, content: str, file_path: Path) -> ToolResult:
        try:
            doc = Document()
            
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

            doc.save(file_path)
            return ToolResult(output=f"DOCX document generated successfully: {file_path}")

        except Exception as e:
            raise e

    def _generate_data_file(self, content: str, file_path: Path, format: str) -> ToolResult:
        try:
            # Try to parse content as JSON
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                # If content is not JSON, assume it's raw string content for simple files,
                # but for structured formats like xlsx/csv/xml we really need structured data.
                # We'll try to wrap it in a structure or fail if strict.
                # For now, let's assume the agent provides JSON string for data.
                return ToolResult(error="Content must be a valid JSON string for data file generation.")

            if format == "json":
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
            
            elif format == "yaml":
                with open(file_path, 'w', encoding='utf-8') as f:
                    yaml.dump(data, f, default_flow_style=False)
            
            elif format == "xml":
                # Basic XML conversion (requires dict/list structure)
                # Using a simple custom converter or pandas if applicable
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                    xml_content = df.to_xml(index=False)
                elif isinstance(data, dict):
                    # Wrap single dict in list for pandas or handle manually
                    # Simple manual XML for dict
                    xml_content = "<root>\n"
                    for k, v in data.items():
                        xml_content += f"  <{k}>{v}</{k}>\n"
                    xml_content += "</root>"
                else:
                     return ToolResult(error="Data must be a list or dictionary for XML generation.")
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(xml_content)

            elif format == "csv":
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                    df.to_csv(file_path, index=False)
                else:
                    return ToolResult(error="Data must be a list of objects for CSV generation.")

            elif format == "xlsx":
                if isinstance(data, list):
                    df = pd.DataFrame(data)
                    df.to_excel(file_path, index=False)
                elif isinstance(data, dict):
                    # If it's a dict of lists, pandas can handle it
                    df = pd.DataFrame(data)
                    df.to_excel(file_path, index=False)
                else:
                    return ToolResult(error="Data must be a list of objects or dict of lists for Excel generation.")

            return ToolResult(output=f"{format.upper()} file generated successfully: {file_path}")

        except Exception as e:
            raise e
